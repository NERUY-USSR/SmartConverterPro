import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from PIL import Image
import os
from pathlib import Path
import threading
import json
from fpdf import FPDF
class FinalConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Smart Converter Pro")
        self.root.geometry("550x480")  # Немного увеличим высоту для нового поля

        self.settings_file = "settings.json"

        # Словари форматов (оставляем без изменений)
        self.format_map = {
            "Изображения": ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'],
            "Тексты": ['.pdf', '.txt', '.docx'],
            "Видео": ['.mp4', '.avi', '.mkv', '.mov'],
            "Аудио": ['.mp3', '.wav', '.ogg', '.m4a', '.flac']
        }

        self.target_options = {
            "Изображения": ["PNG", "JPG", "WEBP", "BMP", "TIFF"],
            "Тексты": ["PDF", "TXT", "DOCX"],
            "Видео": ["MP4", "AVI", "MKV", "MOV"],  # Добавили MOV
            "Аудио": ["MP3", "WAV", "OGG", "M4A", "FLAC"]  # Добавили FLAC
        }

        # 1. ЗАГРУЗКА НАСТРОЕК (Вызываем до создания интерфейса)
        self.load_settings()

        # 2. ПОЛЕ: ПАПКА-ИСТОЧНИК (ОТКУДА БРАТЬ)
        tk.Label(root, text="Откуда брать файлы:", font=("Arial", 10, "bold")).pack(pady=(15, 0))
        in_frame = tk.Frame(root)
        in_frame.pack(fill="x", padx=40)
        self.input_path_var = tk.StringVar(value=self.saved_input_path)
        tk.Entry(in_frame, textvariable=self.input_path_var).pack(side="left", fill="x", expand=True)
        tk.Button(in_frame, text="...", command=self.browse_input).pack(side="right")

        # 3. ПОЛЕ: ВВОД НАЗВАНИЯ (АВТОДОПОЛНЕНИЕ)
        tk.Label(root, text="Введите название файла:", font=("Arial", 10, "bold")).pack(pady=(15, 5))
        self.entry_var = tk.StringVar()

        self.entry = tk.Entry(root, textvariable=self.entry_var, width=60)
        self.entry.pack(pady=5)
        self.entry_var.trace_add("write", self.update_autocomplete)

        # ДОБАВЬТЕ ЭТУ СТРОКУ: Список будет открываться при клике в поле
        self.entry.bind("<FocusIn>", lambda e: self.update_autocomplete())
        self.listbox = tk.Listbox(root, height=5, width=60)
        self.listbox.bind("<<ListboxSelect>>", self.on_file_select)

        self.info_label = tk.Label(root, text="📁 Выберите файл из списка ниже", fg="gray")
        self.info_label.pack(pady=10)

        # 4. ВЫБОР ФОРМАТА
        tk.Label(root, text="Конвертировать в:").pack()
        self.fmt_combo = ttk.Combobox(root, state="readonly", width=15)
        self.fmt_combo.pack(pady=5)

        # 5. ПОЛЕ: ПАПКА СОХРАНЕНИЯ (КУДА)
        tk.Label(root, text="Папка сохранения:").pack(pady=(20, 5))
        out_frame = tk.Frame(root)
        out_frame.pack(fill="x", padx=40)
        self.selected_output_path = tk.StringVar(value=self.saved_output_path)
        tk.Entry(out_frame, textvariable=self.selected_output_path).pack(side="left", fill="x", expand=True)
        tk.Button(out_frame, text="...", command=self.browse_output).pack(side="right")

        # КНОПКА КОНВЕРТАЦИИ
        tk.Button(root, text="КОНВЕРТИРОВАТЬ", bg="#4CAF50", fg="white",
                  font=("Arial", 11, "bold"), height=2, width=25, command=self.start_conversion).pack(pady=30)
        self._current_category = None
        self.root.after(100, self.update_autocomplete)
        # --- НОВЫЕ МЕТОДЫ ДЛЯ ЗАПОМИНАНИЯ ПУТЕЙ ---

    def load_settings(self):
        """Загрузка путей из JSON"""
        self.saved_input_path = os.getcwd()
        self.saved_output_path = os.path.join(Path.home(), "Desktop", "Converted_Files")
        if os.path.exists(self.settings_file):
            try:
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.saved_input_path = data.get("input", self.saved_input_path)
                    self.saved_output_path = data.get("output", self.saved_output_path)
            except:
                pass

    def save_settings(self):
        """Сохранение текущих путей в JSON"""
        data = {
            "input": self.input_path_var.get(),
            "output": self.selected_output_path.get()
        }
        with open(self.settings_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)

    def browse_input(self):
        """Выбор папки 'Откуда'"""
        folder = filedialog.askdirectory()
        if folder:
            self.input_path_var.set(folder)
            self.save_settings()  # Сразу запоминаем
            self.update_autocomplete()  # Сразу обновляем список файлов

    def browse_output(self):
        """Выбор папки 'Куда'"""
        folder = filedialog.askdirectory()
        if folder:
            self.selected_output_path.set(folder)
            self.save_settings()  # Сразу запоминаем
        folder = filedialog.askdirectory()
        if folder: self.selected_output_path.set(folder)

    def update_autocomplete(self, *args):
        search = self.entry_var.get().lower()
        self.listbox.delete(0, tk.END)

        # БЕРЕМ ПУТЬ ИЗ ВАШЕГО НОВОГО ПОЛЯ (ОТКУДА БРАТЬ ФАЙЛЫ)
        source_dir = self.input_path_var.get()

        if not os.path.exists(source_dir):
            self.listbox.insert(tk.END, "❌ Папка не найдена")
            self.listbox.pack(after=self.entry)
            return

        try:
            # ИЩЕМ ФАЙЛЫ ИМЕННО В ЭТОЙ ПАПКЕ
            files = [f for f in os.listdir(source_dir) if os.path.isfile(os.path.join(source_dir, f))]

            # Фильтруем список (если пусто — показываем все)
            if not search:
                matches = files
            else:
                matches = [f for f in files if search in f.lower()]

            if matches:
                for item in matches:
                    ext = os.path.splitext(item)[1].lower()
                    icon = "📄 "
                    if ext in self.format_map["Изображения"]:
                        icon = "🖼️ "
                    elif ext in self.format_map["Видео"]:
                        icon = "🎥 "

                    self.listbox.insert(tk.END, f"{icon}{item}")
                self.listbox.pack(after=self.entry)
            else:
                self.listbox.insert(tk.END, "❌ Файл не найден")
                self.listbox.pack(after=self.entry)

        except Exception:
            self.listbox.pack_forget()

    def on_file_select(self, event):
        if self.listbox.curselection():
            raw_text = self.listbox.get(self.listbox.curselection())

            if "Файл не найден" in raw_text or "Папка не найдена" in raw_text:
                return

            # 1. Убираем иконку (берем всё после первого пробела)
            selected_name = raw_text.split(" ", 1)[1] if " " in raw_text else raw_text

            # 2. В ПОЛЕ ВВОДА пишем только название (для красоты)
            self.entry_var.set(selected_name)
            self.listbox.pack_forget()

            # 3. ОПРЕДЕЛЯЕМ КАТЕГОРИЮ по расширению выбранного файла
            ext = os.path.splitext(selected_name)[1].lower()
            self._current_category = None  # Сбрасываем старую категорию

            for cat, exts in self.format_map.items():
                if ext in exts:
                    # Убираем текущее расширение из списка вариантов
                    opts = [f for f in self.target_options[cat] if f.lower() not in ext]

                    self.info_label.config(text=f"✅ Тип: {cat}", fg="green")
                    self.fmt_combo['values'] = opts
                    if opts:
                        self.fmt_combo.current(0)

                    self._current_category = cat
                    return
            # Если формат не знаком
            self.info_label.config(text="❓ Неизвестный тип файла", fg="red")
            self.fmt_combo['values'] = []

    def start_conversion(self):
        # 1. Склеиваем путь из папки "Источник" и имени файла из поля
        source_dir = self.input_path_var.get()
        short_name = self.entry_var.get()
        filename = os.path.join(source_dir, short_name)

        target_fmt = self.fmt_combo.get()

        # Теперь проверки (имена и пути) будут работать корректно
        if not os.path.exists(filename):
            messagebox.showerror("Ошибка", "Файл ненайден")
            return

        # 1. Проверка существования файла
        if not os.path.exists(filename):
            messagebox.showerror("Ошибка", "Файл ненайден")
            return

        # 2. Проверка выбора формата
        if not target_fmt:
            messagebox.showwarning("Внимание", "Выберите формат!")
            return

        out_dir = self.selected_output_path.get()
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)

        # Берем только имя файла [0], без расширения
        name = os.path.splitext(os.path.basename(filename))[0]
        out_path = os.path.join(out_dir, f"{name}.{target_fmt.lower()}")

        # --- ЛОГИКА ДЛЯ ИЗОБРАЖЕНИЙ ---
        if self._current_category == "Изображения":
            try:
                img = Image.open(filename)
                if target_fmt in ["JPG", "JPEG"] and img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(out_path)
                messagebox.showinfo("Успех", f"Готово!\nПуть: {out_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка картинки: {str(e)}")

        # --- ЛОГИКА ДЛЯ ТЕКСТОВ ---
        elif self._current_category == "Тексты":
            try:
                content = ""  # Изначально текст пустой
                if filename.lower().endswith(".txt"):
                    with open(filename, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()
                elif filename.lower().endswith(".docx"):
                    from docx import Document
                    doc = Document(filename)
                    # ВОТ ЭТА ИСПРАВЛЕННАЯ СТРОКА: собираем текст из всех абзацев
                    content = "\n".join([p.text for p in doc.paragraphs])

                if not content:
                    messagebox.showwarning("Предупреждение", "Файл пуст!")
                    return

                if target_fmt == "PDF":
                    pdf = FPDF()
                    pdf.add_page()

                    # Прямой путь к стандартному шрифту Windows
                    font_path = "C:/Windows/Fonts/arial.ttf"

                    if os.path.exists(font_path):
                        # Добавляем шрифт с поддержкой Unicode
                        pdf.add_font('ArialRus', '', font_path)
                        pdf.set_font('ArialRus', size=12)
                        clean = content
                    else:
                        # Если вдруг файла нет (например, не Windows)
                        pdf.set_font("Helvetica", size=12)
                        clean = content.encode('latin-1', 'replace').decode('latin-1')

                    pdf.multi_cell(0, 10, text=clean)
                    pdf.output(out_path)
                elif target_fmt == "DOCX":
                    from docx import Document
                    d = Document()
                    d.add_paragraph(content)
                    d.save(out_path)
                elif target_fmt == "TXT":
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                messagebox.showinfo("Успех", f"Текст готов!\nПуть: {out_path}")
            except Exception as e:
                messagebox.showerror("Ошибка текста", str(e))
                # --- ЛОГИКА ДЛЯ АУДИО ---
        elif self._current_category == "Аудио":
            try:
                from moviepy import AudioFileClip
                audio = AudioFileClip(filename)

                # Определяем кодек в зависимости от расширения
                ext = target_fmt.lower()
                codec = None

                if ext == 'm4a':
                    codec = 'aac'  # Стандартный кодек вместо проблемного libfdk_aac
                elif ext == 'mp3':
                    codec = 'libmp3lame'

                # Записываем файл
                audio.write_audiofile(out_path, codec=codec, logger=None)

                audio.close()
                messagebox.showinfo("Успех", f"Аудио готово!\nПуть: {out_path}")

            except Exception as e:
                messagebox.showerror("Ошибка аудио", f"Детали: {str(e)}")
                # --- ЛОГИКА ДЛЯ ВИДЕО ---
        elif self._current_category == "Видео":
            #Создаем небольшое окно уведомления
            loading_window = tk.Toplevel(self.root)
            loading_window.title("Подождите")
            loading_window.geometry("300x100")
            tk.Label(loading_window, text="Производится конвертация...\nПожалуйста, не закрывайте программу.",
                        pady=20).pack()

            #Функция, которая будет работать в фоне
            def run_video_conversion():
                try:
                    import subprocess
                    from moviepy import VideoFileClip
                    import multiprocessing

                    #Пытаемся быстро через ffmpeg
                    cmd = f'ffmpeg -i "{filename}" -c copy -y "{out_path}"'
                    result = subprocess.run(cmd, shell=True, capture_output=True)

                    if result.returncode != 0:
                        # Если не вышло быстро, запускаем MoviePy
                        clip = VideoFileClip(filename)
                        clip.write_videofile(out_path, codec="libx264", preset="ultrafast",
                                                threads=multiprocessing.cpu_count(), logger=None)
                        clip.close()

                    # Закрываем окно загрузки и показываем успех в главном потоке
                    loading_window.destroy()
                    messagebox.showinfo("Успех", "Конвертация видео завершена!")
                except Exception as e:
                    loading_window.destroy()
                    messagebox.showerror("Ошибка видео", f"Детали: {str(e)}")
            #Запускаем конвертацию в отдельном потоке, чтобы интерфейс не завис
            threading.Thread(target=run_video_conversion, daemon=True).start()

if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = FinalConverter(root)
        root.mainloop()
    except KeyboardInterrupt:
        pass